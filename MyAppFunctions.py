# -*- coding: utf-8 -*-
"""
Created on Sat Feb 15 15:47:27 2025

@author: jessi
"""

import io
import pandas as pd
from docx import Document
from xml.etree import ElementTree as ET


def check_studentnumber(df, st_num_to_check):

    # Check if the serial number exists
    if st_num_to_check in df['STUDENT_NUMBER'].astype(str).values:
        
        # Find the row(s) where the serial number matches
        record = df[df['STUDENT_NUMBER'] == str(st_num_to_check)]
        return record.reset_index(drop = True)
    else:
        return None
  
    
# Function to convert DataFrame to Word
def to_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name="Sheet1")
    processed_data = output.getvalue()
    return processed_data

# Function to convert DataFrame to Word
def to_word(df):
    doc = Document()
    doc.add_heading("Data Table", level=1)
    
    # Add Table
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name
    
    # Add rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    output = io.BytesIO()
    doc.save(output)
    processed_data = output.getvalue()
    return processed_data

# Function to convert DataFrame to Text
def to_text(df):
    return df.to_csv(index=False, sep="\t")  # Tab-separated text


def read_word_file(file):
    """
    Reads the content of a Word document and returns it as a string.

    Parameters:
    file: Uploaded Word file (BytesIO object)

    Returns:
    str: Text extracted from the document
    """
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])  # Extract paragraphs
    return text

# Function to extract bullet and numbered lists
def extract_list_items_word(doc_path):
    doc = Document(doc_path)
    list_items = []
    for para in doc.paragraphs:
        if para.style.name.startswith("List") or para.text.strip().startswith(tuple("•-123456789")):
            list_items.append(para.text)
    return list_items

# Function to extract bullets, numbers, and equations
def extract_list_and_equations(doc_path):
    doc = Document(doc_path)
    list_items = []
    equations = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if para.style.name.startswith("List") or text.startswith(tuple("•-123456789")):
            list_items.append(text)
        elif "OMath" in [r.name for r in para._element.findall(".//*")]:  # Detects Word equations
            equations.append(text)

    return list_items, equations


